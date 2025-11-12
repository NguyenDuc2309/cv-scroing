import io
from fastapi import HTTPException
import fitz  # PyMuPDF
from docx import Document


async def extract_text_from_pdf(content: bytes) -> str:
    """
    Extract text from PDF file using PyMuPDF (fitz).
    
    Args:
        content: PDF file content as bytes
        
    Returns:
        Extracted text content
    """
    try:
        pdf_file = io.BytesIO(content)
        pdf_document = fitz.open(stream=pdf_file, filetype="pdf")
        
        # Check if PDF is encrypted
        if pdf_document.is_encrypted:
            pdf_document.close()
            raise HTTPException(
                status_code=400,
                detail="PDF file is encrypted. Please provide an unencrypted PDF."
            )
        
        text_parts = []
        for page_num in range(len(pdf_document)):
            page = pdf_document[page_num]
            text = page.get_text()
            if text and text.strip():
                text_parts.append(text.strip())
        
        pdf_document.close()
        
        if not text_parts:
            raise HTTPException(
                status_code=400,
                detail="Could not extract text from PDF. The file might be empty, corrupted, or contain only images."
            )
        
        return "\n".join(text_parts)
    
    except HTTPException:
        raise
    except Exception as e:
        error_msg = str(e)
        # Provide more helpful error messages
        if "'/Root'" in error_msg or "Root" in error_msg:
            raise HTTPException(
                status_code=400,
                detail="Unable to read PDF structure. The PDF may be corrupted or in an unsupported format. Please try converting it to a different PDF version or use DOCX format instead."
            )
        else:
            raise HTTPException(
                status_code=400,
                detail=f"Error reading PDF file: {error_msg}. Please ensure the PDF is not corrupted and contains readable text."
            )


async def extract_text_from_docx(content: bytes) -> str:
    """
    Extract text from DOCX file.
    
    Args:
        content: DOCX file content as bytes
        
    Returns:
        Extracted text content
    """
    try:
        docx_file = io.BytesIO(content)
        doc = Document(docx_file)
        
        text_parts = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)
        
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text_parts.append(" | ".join(row_text))
        
        if not text_parts:
            raise HTTPException(
                status_code=400,
                detail="Could not extract text from DOCX. The file might be empty or corrupted."
            )
        
        return "\n".join(text_parts)
    
    except Exception as e:
        raise HTTPException(
            status_code=400,
            detail=f"Error reading DOCX file: {str(e)}"
        )


async def extract_text(content: bytes, filename: str) -> str:
    """
    Extract text from uploaded file based on file extension.
    
    Args:
        content: File content as bytes
        filename: Name of the uploaded file
        
    Returns:
        Extracted text content
    """
    filename_lower = filename.lower()
    
    if filename_lower.endswith('.pdf'):
        return await extract_text_from_pdf(content)
    elif filename_lower.endswith('.docx'):
        return await extract_text_from_docx(content)
    else:
        raise HTTPException(
            status_code=400,
            detail=f"Unsupported file format. Only PDF and DOCX files are supported."
        )

